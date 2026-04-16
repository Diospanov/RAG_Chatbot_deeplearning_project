"""Load raw PDF and DOCX files into normalized document objects."""

from __future__ import annotations

from datetime import datetime
import logging
from pathlib import Path
import re
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from rag_chatbot.schemas import SourceDocument
from rag_chatbot.utils import make_document_id, normalize_text


LOGGER = logging.getLogger(__name__)


def load_documents(raw_dir: Path, supported_extensions: tuple[str, ...]) -> list[SourceDocument]:
    """Load every supported document from the raw data directory."""

    if not raw_dir.exists():
        raise FileNotFoundError(f"Raw data directory does not exist: {raw_dir}")

    paths = sorted(
        path for path in raw_dir.iterdir() if path.is_file() and path.suffix.lower() in supported_extensions
    )
    if not paths:
        raise FileNotFoundError(
            f"No supported files found in {raw_dir}. Add PDF or DOCX files before running ingestion."
        )

    documents: list[SourceDocument] = []
    for path in paths:
        try:
            if path.suffix.lower() == ".pdf":
                document = _load_pdf(path)
            elif path.suffix.lower() == ".docx":
                document = _load_docx(path)
            else:
                continue
        except Exception as exc:
            LOGGER.warning("Skipping %s because it could not be parsed: %s", path.name, exc)
            continue

        if not document.text:
            LOGGER.warning("Skipping %s because no text was extracted.", path.name)
            continue

        documents.append(document)

    if not documents:
        raise ValueError("Files were found, but no readable text could be extracted from them.")

    return documents


def _load_pdf(path: Path) -> SourceDocument:
    """Extract text and metadata from a PDF file."""

    reader = PdfReader(str(path))
    metadata = reader.metadata or {}
    page_texts = [(page.extract_text() or "").strip() for page in reader.pages]
    text = normalize_text("\n\n".join(part for part in page_texts if part))

    title = _clean_string(metadata.get("/Title")) or path.stem
    date = _normalize_pdf_date(metadata.get("/CreationDate") or metadata.get("/ModDate"))

    normalized_metadata = {
        "author": _clean_string(metadata.get("/Author")),
        "subject": _clean_string(metadata.get("/Subject")),
        "producer": _clean_string(metadata.get("/Producer")),
        "creator": _clean_string(metadata.get("/Creator")),
        "page_count": len(reader.pages),
    }

    return SourceDocument(
        document_id=make_document_id(path.name),
        source_path=str(path.resolve()),
        source_filename=path.name,
        source_type="pdf",
        title=title,
        date=date,
        text=text,
        char_count=len(text),
        metadata={key: value for key, value in normalized_metadata.items() if value not in (None, "")},
    )


def _load_docx(path: Path) -> SourceDocument:
    """Extract text and metadata from a DOCX file."""

    document = DocxDocument(str(path))
    properties = document.core_properties
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    text = normalize_text("\n\n".join(paragraphs))

    title = _clean_string(properties.title) or path.stem
    date = _normalize_datetime(properties.created or properties.modified)

    normalized_metadata = {
        "author": _clean_string(properties.author),
        "category": _clean_string(properties.category),
        "comments": _clean_string(properties.comments),
        "subject": _clean_string(properties.subject),
        "last_modified_by": _clean_string(properties.last_modified_by),
        "paragraph_count": len(paragraphs),
    }

    return SourceDocument(
        document_id=make_document_id(path.name),
        source_path=str(path.resolve()),
        source_filename=path.name,
        source_type="docx",
        title=title,
        date=date,
        text=text,
        char_count=len(text),
        metadata={key: value for key, value in normalized_metadata.items() if value not in (None, "")},
    )


def _clean_string(value: Any) -> str | None:
    """Normalize metadata text fields into stripped strings."""

    if value is None:
        return None
    cleaned = str(value).strip()
    return cleaned or None


def _normalize_datetime(value: datetime | None) -> str | None:
    """Convert Python datetimes to ISO format for JSON output."""

    if value is None:
        return None
    return value.isoformat()


def _normalize_pdf_date(value: Any) -> str | None:
    """Parse common PDF metadata date values when available."""

    cleaned = _clean_string(value)
    if not cleaned:
        return None

    match = re.match(r"D:(\d{4})(\d{2})?(\d{2})?(\d{2})?(\d{2})?(\d{2})?", cleaned)
    if not match:
        return cleaned

    parts = match.groups(default="01")
    try:
        parsed = datetime(
            int(parts[0]),
            int(parts[1]),
            int(parts[2]),
            int(parts[3]),
            int(parts[4]),
            int(parts[5]),
        )
    except ValueError:
        return cleaned
    return parsed.isoformat()
